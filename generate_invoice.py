"""
Tax Invoice generator (SLNS LOGISTICS style).

PHASE 1: `load_invoice_data()` - all invoice values hardcoded, matching the
sample invoice image. Kept around for reference/testing without an Excel
file on hand.

PHASE 2 (current default): `load_invoice_data_from_excel()` reads the same
data out of Invoice_Data.xlsx (see create_excel_template.py for the sheet
layout - label in column A, value in column B, blank rows are just visual
section breaks). GST/CGST/SGST and every total are computed here, never
read from the sheet. Either loader returns the same dict shape, so
`build_invoice()` itself never needs to change.
"""

import os
import datetime

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import load_workbook

DEFAULT_FONT = "Arial"

# Column-A label text for every dynamic field in Invoice_Data.xlsx. Shared between
# the reader below and create_excel_template.py so the two can never drift apart.
EXCEL_FIELD_LABELS = {
    "invoice_no": "Invoice No",
    "invoice_date": "Invoice Date",
    "due_date": "Due Date",
    "place_of_supply": "Place of Supply",
    "bill_to_name": "Bill To Company Name",
    "bill_to_address": "Bill To Address",
    "bill_to_state": "Bill To State",
    "bill_to_gstin": "Bill To GSTIN",
    "shipping_address": "Shipping Address",
    "transport_name": "Transport Name",
    "transport_delivery_date": "Delivery Date",
    "transport_vehicle_number": "Vehicle Number",
    "transport_delivery_location": "Delivery Location",
    "item_name": "Item Name",
    "item_quantity": "Quantity",
    "item_unit": "Unit",
    "item_price_per_unit": "Price Per Unit",
    "amount_in_words": "Invoice Amount In Words",
    "bank_name": "Bank Name",
    "bank_account_no": "Account No",
    "bank_ifsc": "IFSC Code",
    "bank_account_holder": "Account Holder Name",
    "signature": "Signature",
}


# ---------------------------------------------------------------------------
# PHASE 1 DATA — replace this block with an Excel-reader in phase 2.
# Every field here corresponds 1:1 to a label in the source invoice image.
# ---------------------------------------------------------------------------
def load_invoice_data():
    return {
        "company": {
            "name": "SLNS LOGISTICS",
            "address_lines": [
                "Plot No : 118, Balaji Nagar, Beeramguda,",
                "Hyderabad, Sangareddy, Telangana - 502032",
            ],
            "phone": "6304066403",
            "email": "Logisticsslns@gmail.com",
            "gstin": "36ANLPK6091R1Z2",
            "state": "36-Telangana",
        },
        "invoice": {
            "no": "INV-214",
            "date": "24-Jun-2026",
            "due_date": "26-Jun-2026",
            "place_of_supply": "36-Telangana",
        },
        "bill_to": {
            "name": "HEARTCULTURE NATURAL PRODUCT LLP",
            "address_lines": [
                "13-110 Kanha village nandigama mandal",
                "Ranga Reddy telangana 509325",
            ],
            "state": "36-Telangana",
            "gstin": "36AAMFH4243C1ZL",
        },
        "shipping_address": "Rajasthan",
        "transport": {
            "name": "-",
            "delivery_date": "-",
            "vehicle_number": "RJ 19 GF 8254",
            "delivery_location": "Rajasthan",
        },
        "items": [
            {
                "sno": 1,
                "name": "Transport charges for Wakavali to khana hyderabad",
                "qty": "10",
                "unit": "Ton",
                "price_per_unit": "4,900.00",
                "gst_amount": "2,450.00",
                "gst_rate": "5.0",
                "amount": "51,450.00",
            }
        ],
        "tax_summary": {
            "taxable_amount": "49,000.00",
            "cgst_rate": "2.5",
            "cgst_amt": "1,225.00",
            "sgst_rate": "2.5",
            "sgst_amt": "1,225.00",
            "total_tax": "2,450.00",
        },
        "totals": {
            "sub_total": "51,450.00",
            "total_tax": "2,450.00",
            "total": "51,450.00",
        },
        "amount_in_words": "Fifty One Thousand Four Hundred and Fifty Rupees only",
        "received": "0.00",
        "balance": "51,450.00",
        "bank": {
            "name": "Hdfc Bank, Serilingampally",
            "account_no": "50200118563217",
            "ifsc": "HDFC0002073",
            "account_holder": "SLNS LOGISTICS",
        },
        "signatory": {
            "for_company": "For SLNS LOGISTICS",
            "signatory_name": "Bharathil",
            "label": "Authorized Signatory",
        },
    }


def _num(value, default=0.0):
    """Coerce an Excel cell value to float, tolerating None/blank cells."""
    if value is None or value == "":
        return default
    return float(value)


def _fmt_value(value):
    """Render an Excel cell value as display text - Excel may hand back a
    real datetime even when the sheet displays a plain date string."""
    if isinstance(value, (datetime.datetime, datetime.date)):
        return value.strftime("%d-%b-%Y")
    if value is None:
        return ""
    return str(value).strip()


def _fmt_money(value):
    return f"{value:,.2f}"


def _fmt_qty(value):
    return str(int(value)) if float(value).is_integer() else str(value)


def load_invoice_data_from_excel(path, image_out_dir=None):
    """Phase 2 loader: reads Invoice_Data.xlsx (label in column A, value in
    column B - blank rows are just visual section breaks in the sheet, not
    meaningful to this code, which matches by label text instead) and
    returns the same dict shape load_invoice_data() does.

    GST/CGST/SGST and every total are computed here (5% GST, split 2.5%
    CGST / 2.5% SGST off the taxable amount) - none of that is read from the
    sheet. Received is always 0 and Balance always equals the grand total.
    The signature is whatever picture is embedded in the sheet (anywhere) -
    extracted to a PNG next to the workbook and inserted as an image instead
    of typed text.
    """
    wb = load_workbook(path, data_only=True)
    ws = wb.active

    raw = {}
    for row in ws.iter_rows():
        label_cell = row[0]
        value_cell = row[1] if len(row) > 1 else None
        label = label_cell.value
        if label is None or str(label).strip() == "":
            continue
        raw[str(label).strip()] = value_cell.value if value_cell is not None else None

    def get(key):
        return raw.get(EXCEL_FIELD_LABELS[key])

    # ---- Signature: first picture embedded anywhere in the sheet ----
    image_out_dir = image_out_dir or os.path.dirname(os.path.abspath(path))
    signature_path = None
    images = getattr(ws, "_images", [])
    if images:
        signature_path = os.path.join(image_out_dir, "_signature_extracted.png")
        with open(signature_path, "wb") as f:
            f.write(images[0]._data())

    # ---- Line item + GST math (5% GST, split 2.5% CGST / 2.5% SGST) ----
    qty = _num(get("item_quantity"))
    price_per_unit = _num(get("item_price_per_unit"))
    taxable_amount = qty * price_per_unit
    gst_rate = 5.0
    cgst_rate = gst_rate / 2
    sgst_rate = gst_rate / 2
    cgst_amt = taxable_amount * cgst_rate / 100
    sgst_amt = taxable_amount * sgst_rate / 100
    gst_amount = cgst_amt + sgst_amt
    line_amount = taxable_amount + gst_amount
    grand_total = line_amount

    bill_to_address_lines = _fmt_value(get("bill_to_address")).split("\n")

    return {
        "company": {
            "name": "SLNS LOGISTICS",
            "address_lines": [
                "Plot No : 118, Balaji Nagar, Beeramguda,",
                "Hyderabad, Sangareddy, Telangana - 502032",
            ],
            "phone": "6304066403",
            "email": "Logisticsslns@gmail.com",
            "gstin": "36ANLPK6091R1Z2",
            "state": "36-Telangana",
        },
        "invoice": {
            "no": _fmt_value(get("invoice_no")),
            "date": _fmt_value(get("invoice_date")),
            "due_date": _fmt_value(get("due_date")),
            "place_of_supply": _fmt_value(get("place_of_supply")),
        },
        "bill_to": {
            "name": _fmt_value(get("bill_to_name")),
            "address_lines": bill_to_address_lines,
            "state": _fmt_value(get("bill_to_state")),
            "gstin": _fmt_value(get("bill_to_gstin")),
        },
        "shipping_address": _fmt_value(get("shipping_address")),
        "transport": {
            "name": _fmt_value(get("transport_name")),
            "delivery_date": _fmt_value(get("transport_delivery_date")),
            "vehicle_number": _fmt_value(get("transport_vehicle_number")),
            "delivery_location": _fmt_value(get("transport_delivery_location")),
        },
        "items": [
            {
                "sno": 1,
                "name": _fmt_value(get("item_name")),
                "qty": _fmt_qty(qty),
                "unit": _fmt_value(get("item_unit")),
                "price_per_unit": _fmt_money(price_per_unit),
                "gst_amount": _fmt_money(gst_amount),
                "gst_rate": f"{gst_rate:g}",
                "amount": _fmt_money(line_amount),
            }
        ],
        "tax_summary": {
            "taxable_amount": _fmt_money(taxable_amount),
            "cgst_rate": f"{cgst_rate:g}",
            "cgst_amt": _fmt_money(cgst_amt),
            "sgst_rate": f"{sgst_rate:g}",
            "sgst_amt": _fmt_money(sgst_amt),
            "total_tax": _fmt_money(gst_amount),
        },
        "totals": {
            "sub_total": _fmt_money(line_amount),
            "total_tax": _fmt_money(gst_amount),
            "total": _fmt_money(grand_total),
        },
        "amount_in_words": _fmt_value(get("amount_in_words")),
        "received": "0.00",
        "balance": _fmt_money(grand_total),
        "bank": {
            "name": _fmt_value(get("bank_name")),
            "account_no": _fmt_value(get("bank_account_no")),
            "ifsc": _fmt_value(get("bank_ifsc")),
            "account_holder": _fmt_value(get("bank_account_holder")),
        },
        "signatory": {
            "for_company": "For SLNS LOGISTICS",
            "signatory_name": "",
            "label": "Authorized Signatory",
            "image_path": signature_path,
        },
    }


# ---------------------------------------------------------------------------
# Low level helpers
# ---------------------------------------------------------------------------
def set_cell_border(cell, **kwargs):
    """Set individual borders on a table cell. kwargs keys: top, bottom, start, end."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "bottom", "start", "end", "left", "right"):
        edge_data = kwargs.get(edge)
        if edge_data is None:
            continue
        tag = f"w:{edge}"
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def all_borders(cell, sz=4, color="000000", val="single"):
    """Set top/bottom/left/right on a cell. Both start/end (logical) and left/right
    (physical) are set to the same values - some Word versions only honor one pair,
    and relying on just start/end can leave one edge of a table borderless."""
    edge = {"sz": sz, "val": val, "color": color}
    set_cell_border(cell, top=edge, bottom=edge, start=edge, end=edge, left=edge, right=edge)


def set_run_font(run, name=DEFAULT_FONT):
    """Force a font on a run, overriding theme fallback for both western and east-asian/cs ranges."""
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)


def set_cell_text(cell, text, bold=False, size=8, align=WD_ALIGN_PARAGRAPH.LEFT,
                   color=None, vertical_center=True, font=DEFAULT_FONT):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    set_run_font(run, font)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if vertical_center:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return run


def add_cell_lines(cell, lines, vertical_center=False, font=DEFAULT_FONT):
    """lines: list where each item is one paragraph.
    An item is either a (text, bold, size) tuple (single run), or a list of
    such tuples rendered as multiple runs on the same line (e.g. a bold
    label followed by a plain value)."""
    cell.text = ""
    first = True
    for line in lines:
        para = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        segments = line if isinstance(line, list) else [line]
        for text, bold, size in segments:
            run = para.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            set_run_font(run, font)
    if vertical_center:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_tab(para):
    """Insert a real tab character (<w:tab/>) so text after it jumps to the paragraph's tab stop."""
    run = para.add_run()
    run._element.append(OxmlElement("w:tab"))
    return run


def add_label_line(container, label, value, tab_cm, size=8, colon=": ",
                    font=DEFAULT_FONT, first=False, label_bold=False, value_tab_cm=None):
    """One 'Label <tab> : value' line, with the tab stop making the colon column line up
    across every line in the block regardless of label length (labels are proportional-font,
    so fixed spaces don't align them).

    If value_tab_cm is given, a second tab stop is added and the value starts there -
    i.e. a real tab character after the colon, rather than just a space - with every
    line in the block still landing on the same value column."""
    para = container.paragraphs[0] if first else container.add_paragraph()
    para.paragraph_format.tab_stops.add_tab_stop(Cm(tab_cm), WD_TAB_ALIGNMENT.LEFT)
    r_label = para.add_run(label)
    r_label.bold = label_bold
    r_label.font.size = Pt(size)
    set_run_font(r_label, font)
    add_tab(para)
    if value_tab_cm is not None:
        para.paragraph_format.tab_stops.add_tab_stop(Cm(value_tab_cm), WD_TAB_ALIGNMENT.LEFT)
        r_colon = para.add_run(colon.strip())
        r_colon.font.size = Pt(size)
        set_run_font(r_colon, font)
        add_tab(para)
        r_value = para.add_run(str(value))
    else:
        r_value = para.add_run(f"{colon}{value}")
    r_value.font.size = Pt(size)
    set_run_font(r_value, font)
    return para


def add_dual_label_line(container, label1, value1, label2, value2, tab_cm,
                         size=8, colon=": ", font=DEFAULT_FONT, first=False):
    """One line with two 'Label: value' pairs, where the second pair always starts at
    the same tab-stop position regardless of how long the first value is."""
    para = container.paragraphs[0] if first else container.add_paragraph()
    para.paragraph_format.tab_stops.add_tab_stop(Cm(tab_cm), WD_TAB_ALIGNMENT.LEFT)
    r1 = para.add_run(label1)
    r1.bold = True
    r1.font.size = Pt(size)
    set_run_font(r1, font)
    r2 = para.add_run(f"{colon}{value1}")
    r2.font.size = Pt(size)
    set_run_font(r2, font)
    add_tab(para)
    r3 = para.add_run(label2)
    r3.bold = True
    r3.font.size = Pt(size)
    set_run_font(r3, font)
    r4 = para.add_run(f"{colon}{value2}")
    r4.font.size = Pt(size)
    set_run_font(r4, font)
    return para


def label_value_row(table, row_idx, label, value, label_size=8, value_size=8):
    label_cell = table.cell(row_idx, 0)
    colon_cell = table.cell(row_idx, 1)
    value_cell = table.cell(row_idx, 2)
    set_cell_text(label_cell, label, bold=False, size=label_size)
    set_cell_text(colon_cell, ":", bold=False, size=label_size, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(value_cell, str(value), bold=False, size=value_size)


def set_col_widths(table, widths):
    widths = [Emu(int(w)) for w in widths]
    table.autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w
    for col, w in zip(table.columns, widths):
        col.width = w


def shade_cell(cell, fill="D9D9D9"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def apply_borders_to_table(table, sz=4):
    for row in table.rows:
        for cell in row.cells:
            all_borders(cell, sz=sz)


def add_gap(doc, pt=2):
    """Small vertical spacer between tables (docx tables can't be adjacent with custom gaps)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("")
    run.font.size = Pt(pt)
    set_run_font(run)


def set_default_font(doc, name=DEFAULT_FONT):
    """Set the theme/style default so any run without an explicit font still renders as `name`."""
    style = doc.styles["Normal"]
    style.font.name = name
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------
def build_invoice(data, out_path="Tax_Invoice.docx"):
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)

    set_default_font(doc)
    normal = doc.styles["Normal"]
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0

    # ---- Title ----
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TAX INVOICE")
    run.bold = True
    run.font.size = Pt(18)
    set_run_font(run)

    # ---- Row 1: Company details | Invoice details ----
    # Real usable width (page width minus both margins) - not a guessed constant, so every
    # full-width table below lines up on the same left/right edges as the page margins.
    total_width = Emu(int(section.page_width) - int(section.left_margin) - int(section.right_margin))

    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_col_widths(header_table, [total_width * 0.6, total_width * 0.4])

    company = data["company"]
    left_cell = header_table.cell(0, 0)
    add_cell_lines(left_cell, [
        (company["name"], True, 16),
        *[(line, False, 8) for line in company["address_lines"]],
        ("", False, 2),
    ])
    # Phone/Email and GSTIN/State share one tab stop so "Email" and "State" always
    # start at the same horizontal position, no matter how long the phone/GSTIN value is.
    COMPANY_TAB_CM = 5.1
    add_dual_label_line(left_cell, "Phone", company["phone"], "Email", company["email"],
                        tab_cm=COMPANY_TAB_CM)
    add_dual_label_line(left_cell, "GSTIN", company["gstin"], "State", company["state"],
                        tab_cm=COMPANY_TAB_CM)

    invoice = data["invoice"]
    right_cell = header_table.cell(0, 1)
    right_cell.text = ""
    p = right_cell.paragraphs[0]
    r = p.add_run("Invoice Details:")
    r.bold = True
    r.font.size = Pt(8)
    set_run_font(r)
    INVOICE_TAB_CM = 3.4
    INVOICE_VALUE_TAB_CM = INVOICE_TAB_CM + 1.25
    for label, value in [
        ("No", invoice["no"]),
        ("Date", invoice["date"]),
        ("Due Date", invoice["due_date"]),
        ("Place of Supply", invoice["place_of_supply"]),
    ]:
        add_label_line(right_cell, label, value, tab_cm=INVOICE_TAB_CM,
                       value_tab_cm=INVOICE_VALUE_TAB_CM)

    apply_borders_to_table(header_table)
    for cell in header_table.rows[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # ---- Row 2: Bill To | Shipping Address | Transportation Details ----
    add_gap(doc)
    tri_table = doc.add_table(rows=1, cols=3)
    tri_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Bill To 40% / Shipping Address 20% / Transportation Details 40%.
    set_col_widths(tri_table, [total_width * 0.4, total_width * 0.2, total_width * 0.4])

    bill_to = data["bill_to"]
    bill_cell = tri_table.cell(0, 0)
    add_cell_lines(bill_cell, [
        ("Bill To:", True, 8),
        (bill_to["name"], True, 8),
        *[(line, False, 8) for line in bill_to["address_lines"]],
        ("", False, 2),
    ])
    BILL_TO_TAB_CM = 3.1
    add_label_line(bill_cell, "State", bill_to["state"], tab_cm=BILL_TO_TAB_CM, label_bold=True)
    add_label_line(bill_cell, "GSTIN Number", bill_to["gstin"], tab_cm=BILL_TO_TAB_CM, label_bold=True)

    ship_cell = tri_table.cell(0, 1)
    add_cell_lines(ship_cell, [
        ("Shipping Address:", True, 8),
        (data["shipping_address"], False, 8),
    ])

    transport = data["transport"]
    transport_cell = tri_table.cell(0, 2)
    transport_cell.text = ""
    p = transport_cell.paragraphs[0]
    r = p.add_run("Transportation Details:")
    r.bold = True
    r.font.size = Pt(8)
    set_run_font(r)
    TRANSPORT_TAB_CM = 3.8
    TRANSPORT_VALUE_TAB_CM = TRANSPORT_TAB_CM + 1.25
    for label, value in [
        ("Transport Name", transport["name"]),
        ("Delivery Date", transport["delivery_date"]),
        ("Vehicle Number", transport["vehicle_number"]),
        ("Delivery location", transport["delivery_location"]),
    ]:
        add_label_line(transport_cell, label, value, tab_cm=TRANSPORT_TAB_CM,
                       value_tab_cm=TRANSPORT_VALUE_TAB_CM)

    apply_borders_to_table(tri_table)
    for cell in tri_table.rows[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # ---- Items table ----
    add_gap(doc)
    headers = ["#", "Item Name", "Quantity", "Unit", "Price/ Unit (Rs.)", "GST(Rs.)", "Amount(Rs.)"]
    items = data["items"]
    items_table = doc.add_table(rows=1 + len(items) + 1, cols=len(headers))
    items_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Column proportions, scaled to the page's real usable width so the table's right
    # edge always lands exactly on the margin instead of stopping short of it.
    items_col_ratios = [0.9, 6.0, 2.3, 1.8, 3.0, 2.5, 2.5]
    ratio_sum = sum(items_col_ratios)
    widths = [total_width * (c / ratio_sum) for c in items_col_ratios]
    set_col_widths(items_table, widths)

    for idx, h in enumerate(headers):
        cell = items_table.cell(0, idx)
        set_cell_text(cell, h, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(cell)

    for r, item in enumerate(items, start=1):
        row_vals = [
            str(item["sno"]), item["name"], item["qty"], item["unit"],
            item["price_per_unit"], f"{item['gst_amount']}\n({item['gst_rate']}%)",
            item["amount"],
        ]
        for c, val in enumerate(row_vals):
            align = WD_ALIGN_PARAGRAPH.CENTER if c != 1 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(items_table.cell(r, c), val, align=align)

    # Total row (merged label across first 4 cols)
    total_row_idx = len(items) + 1
    total_row = items_table.rows[total_row_idx]
    a = items_table.cell(total_row_idx, 0)
    b = items_table.cell(total_row_idx, 3)
    a.merge(b)
    set_cell_text(a, "Total", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(items_table.cell(total_row_idx, 4), "", bold=True)
    set_cell_text(items_table.cell(total_row_idx, 5), data["totals"]["total_tax"], bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(items_table.cell(total_row_idx, 6), data["totals"]["sub_total"], bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    apply_borders_to_table(items_table)

    add_gap(doc)

    # ---- Tax Summary ----
    tax_label_run = doc.add_paragraph().add_run("Tax Summary:")
    tax_label_run.bold = True
    tax_label_run.font.size = Pt(8)
    set_run_font(tax_label_run)

    # Columns: Taxable Amount | CGST Rate | CGST Amt | SGST Rate | SGST Amt | Total Tax
    # Same total_width-ratio approach as the items table, so this table's edges also
    # land exactly on the page margins instead of stopping short of them.
    tax = data["tax_summary"]
    tax_table = doc.add_table(rows=4, cols=6)
    tax_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tax_col_ratios = [4.0, 2.8, 3.1, 2.8, 3.1, 3.1]
    tax_ratio_sum = sum(tax_col_ratios)
    set_col_widths(tax_table, [total_width * (c / tax_ratio_sum) for c in tax_col_ratios])

    # Row 0/1: header (Taxable Amount + Total Tax span both header rows; CGST/SGST span 2 cols)
    tax_table.cell(0, 0).merge(tax_table.cell(1, 0))
    tax_table.cell(0, 1).merge(tax_table.cell(0, 2))
    tax_table.cell(0, 3).merge(tax_table.cell(0, 4))
    tax_table.cell(0, 5).merge(tax_table.cell(1, 5))

    set_cell_text(tax_table.cell(0, 0), "Taxable Amount (Rs.)", bold=True, size=7,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(tax_table.cell(0, 1), "CGST", bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(tax_table.cell(0, 3), "SGST", bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(tax_table.cell(0, 5), "Total Tax(Rs.)", bold=True, size=7,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(tax_table.cell(1, 1), "Rate (%)", bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(tax_table.cell(1, 2), "Amt (Rs.)", bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(tax_table.cell(1, 3), "Rate (%)", bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(tax_table.cell(1, 4), "Amt (Rs.)", bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 2: data row
    for c, val in enumerate([tax["taxable_amount"], tax["cgst_rate"], tax["cgst_amt"],
                              tax["sgst_rate"], tax["sgst_amt"], tax["total_tax"]]):
        set_cell_text(tax_table.cell(2, c), str(val), size=7, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 3: TOTAL row (rate columns blank)
    set_cell_text(tax_table.cell(3, 0), "TOTAL", bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(tax_table.cell(3, 1), "", size=7)
    set_cell_text(tax_table.cell(3, 2), tax["cgst_amt"], bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(tax_table.cell(3, 3), "", size=7)
    set_cell_text(tax_table.cell(3, 4), tax["sgst_amt"], bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(tax_table.cell(3, 5), tax["total_tax"], bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER)

    apply_borders_to_table(tax_table)

    add_gap(doc)

    # ---- Amount in words ----
    words_table = doc.add_table(rows=1, cols=1)
    words_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_col_widths(words_table, [total_width])
    set_cell_text(words_table.cell(0, 0),
                  f"Invoice Amount In Words :   {data['amount_in_words']}", bold=True, size=8)
    apply_borders_to_table(words_table)

    # ---- Received / Balance ----
    rb_table = doc.add_table(rows=1, cols=2)
    rb_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_col_widths(rb_table, [total_width * 0.5, total_width * 0.5])
    set_cell_text(rb_table.cell(0, 0), f"Received :  Rs. {data['received']}", bold=True)
    set_cell_text(rb_table.cell(0, 1), f"Balance :  Rs. {data['balance']}", bold=True,
                  align=WD_ALIGN_PARAGRAPH.RIGHT)
    apply_borders_to_table(rb_table)

    add_gap(doc)

    # ---- Bank details + Signatory ----
    # A standalone table, one blank line below the Received/Balance row.
    footer_table = doc.add_table(rows=1, cols=2)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_col_widths(footer_table, [total_width * 0.6, total_width * 0.4])
    bank = data["bank"]
    bank_cell = footer_table.cell(0, 0)
    bank_cell.text = ""
    p = bank_cell.paragraphs[0]
    r = p.add_run("Bank Details:")
    r.bold = True
    r.font.size = Pt(9)
    set_run_font(r)
    BANK_TAB_CM = 4.6
    for label, value in [
        ("Name", bank["name"]),
        ("Account No.", bank["account_no"]),
        ("IFSC code", bank["ifsc"]),
        ("Account Holder's Name", bank["account_holder"]),
    ]:
        add_label_line(bank_cell, label, value, tab_cm=BANK_TAB_CM)

    sign = data["signatory"]
    sign_cell = footer_table.cell(0, 1)
    sign_cell.text = ""
    p0 = sign_cell.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(sign["for_company"])
    r0.bold = True
    r0.font.size = Pt(8)
    set_run_font(r0)

    image_path = sign.get("image_path")
    if image_path and os.path.exists(image_path):
        # Actual signature picture (extracted from the Excel sheet) instead of typed text.
        p_img = sign_cell.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(4)
        p_img.paragraph_format.space_after = Pt(0)
        p_img.add_run().add_picture(image_path, height=Cm(1.0))
    else:
        for _ in range(2):
            gap_p = sign_cell.add_paragraph()
            gap_p.paragraph_format.space_before = Pt(0)
            gap_p.paragraph_format.space_after = Pt(0)
            gap_run = gap_p.add_run("")
            gap_run.font.size = Pt(2)
            set_run_font(gap_run)
        name_p = sign_cell.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_p.add_run(sign["signatory_name"])
        name_run.font.size = Pt(10)
        set_run_font(name_run)

    label_p = sign_cell.add_paragraph()
    label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = label_p.add_run(sign["label"])
    label_run.font.size = Pt(8)
    set_run_font(label_run)

    apply_borders_to_table(footer_table)

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    excel_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Invoice_Data.xlsx")
    if os.path.exists(excel_path):
        data = load_invoice_data_from_excel(excel_path)
    else:
        print(f"{excel_path} not found - falling back to the hardcoded phase 1 sample data.")
        data = load_invoice_data()
    path = build_invoice(data, out_path="Tax_Invoice.docx")
    print(f"Invoice generated: {path}")
